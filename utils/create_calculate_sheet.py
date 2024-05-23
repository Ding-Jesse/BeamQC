import docx
from docx.shared import Pt, RGBColor
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from copy import deepcopy
from typing import Union
from item.column import Column

from PIL import Image
from docx.shared import Inches, Cm

# Function to add an OMML equation to a paragraph


def add_omml_equation(paragraph, parts):
    run = paragraph.add_run()
    omml = OxmlElement('m:oMathPara')
    math = OxmlElement('m:oMath')
    for part in parts:
        if isinstance(part, str):
            part = create_math_element(part)
        math.append(part)
    # math.append(equation)
    omml.append(math)
    run._r.append(omml)
# Function to create a superscript element


def create_superscript(base, superscript):
    sSup = OxmlElement('m:sSup')
    e = OxmlElement('m:e')
    e_r = OxmlElement('m:r')
    e_t = OxmlElement('m:t')
    e_t.text = base
    e_r.append(e_t)
    e.append(e_r)

    sup = OxmlElement('m:sup')
    sup_r = OxmlElement('m:r')
    sup_t = OxmlElement('m:t')
    sup_t.text = superscript
    sup_r.append(sup_t)
    sup.append(sup_r)

    sSup.append(e)
    sSup.append(sup)
    return sSup

# Function to create a subscript element


def create_subscript(base, subscript):
    sSub = OxmlElement('m:sSub')
    e = OxmlElement('m:e')
    e_r = OxmlElement('m:r')
    e_t = OxmlElement('m:t')
    e_t.text = base
    e_r.append(e_t)
    e.append(e_r)

    sub = OxmlElement('m:sub')
    sub_r = OxmlElement('m:r')
    sub_t = OxmlElement('m:t')
    sub_t.text = subscript
    sub_r.append(sub_t)
    sub.append(sub_r)

    sSub.append(e)
    sSub.append(sub)
    return sSub


def create_superscript_subscript(base, superscript, subscript):
    sSubSup = OxmlElement('m:sSubSup')

    # Base element
    e = OxmlElement('m:e')
    e_r = OxmlElement('m:r')
    e_t = OxmlElement('m:t')
    e_t.text = base
    e_r.append(e_t)
    e.append(e_r)

    # Superscript element
    sup = OxmlElement('m:sup')
    sup_r = OxmlElement('m:r')
    sup_t = OxmlElement('m:t')
    sup_t.text = superscript
    sup_r.append(sup_t)
    sup.append(sup_r)

    # Subscript element
    sub = OxmlElement('m:sub')
    sub_r = OxmlElement('m:r')
    sub_t = OxmlElement('m:t')
    sub_t.text = subscript
    sub_r.append(sub_t)
    sub.append(sub_r)

    sSubSup.append(e)
    sSubSup.append(sub)
    sSubSup.append(sup)
    return sSubSup


def create_fraction(numerators, denominators):
    f = OxmlElement('m:f')

    # Numerator element
    num = OxmlElement('m:num')
    # num_r = OxmlElement('m:r')
    # num_t = OxmlElement('m:t')
    # num_t.text = numerator
    # num_r.append(num_t)
    for numerator in numerators:
        num.append(numerator)

    # Denominator element
    den = OxmlElement('m:den')
    # den_r = OxmlElement('m:r')
    # den_t = OxmlElement('m:t')
    # den_t.text = denominator
    # den_r.append(den_t)
    for denominator in denominators:
        den.append(denominator)

    f.append(num)
    f.append(den)
    return f


def create_math_element(text):
    r = OxmlElement('m:r')
    t = OxmlElement('m:t')
    t.text = text
    r.append(t)
    return r


def add_image_to_doc(doc: DocumentObject, image_path):
    # Add a paragraph for the image
    paragraph = doc.add_paragraph()
    # Add an image with its original size
    run = paragraph.add_run()
    run.add_picture(image_path)

    # Center-align the paragraph
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_square_root(element):
    if isinstance(element, str):
        element = create_math_element(element)
    rad = OxmlElement('m:rad')
    # deg = OxmlElement('m:deg')
    e = OxmlElement('m:e')

    # rad.append(deg)
    e.append(element)
    rad.append(e)

    return rad


def set_chinese_font(run: Union[Run, Paragraph], font: str):
    if isinstance(run, Paragraph):
        run = run.runs[0]
    # Ensure the font is applied correctly by setting the eastAsian property
    rPr = run._r.get_or_add_rPr()
    eastAsia_font = OxmlElement('w:rFonts')
    eastAsia_font.set(qn('w:eastAsia'), font)
    rPr.append(eastAsia_font)

# Apply border to the table


def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))
            tcPr.append(element)

# Create a new Document


def get_page_width(section):
    """
    Get the width of the page from the section properties.
    """
    # Retrieve the width of the page in twips (twentieths of a point)
    page_width_twips = section.page_width
    # Convert twips to inches (1 inch = 1440 twips)
    page_width_inches = page_width_twips / 1440
    return page_width_inches


def add_scaled_image(doc: DocumentObject, image_path):
    """
    Add an image to the document scaled to fit the page width.
    """
    # Open the image to get its dimensions
    page_width_inches = get_page_width(doc.sections[0])
    with Image.open(image_path) as img:
        img_width, img_height = img.size

    # Calculate the width and height in inches, maintaining the aspect ratio
    img_aspect_ratio = img_height / img_width
    max_width = page_width_inches - 1  # Leave some margin on the sides
    new_width = min(max_width, img_width / img.info.get('dpi', (96, 96))[0])
    new_height = new_width * img_aspect_ratio

    # Add a paragraph for the image
    paragraph = doc.add_paragraph()
    # Add an image with its scaled size
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(new_width))
    # doc.add_picture(image_path, width=Cm(new_width))
    # Center-align the paragraph
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_calculate_sheet(doc_filename, column_list: list[Column], output_serial: list = [], output_floor: list = []):

    doc = docx.Document()
    chinese_font = '微軟正黑體'
    style = doc.styles['Normal']
    font = style.font
    # font.name = '微軟正黑體'
    # font.eastAsia = '微軟正黑體'
    font.size = Pt(12)

    # Add a title
    p = doc.add_heading('耐震接頭檢討', level=1)
    set_chinese_font(p, font=chinese_font)
    # Add a section for "接頭剪力設計原則"
    p = doc.add_heading('接頭剪力設計原則', level=2)
    set_chinese_font(p, font=chinese_font)

    add_image_to_doc(doc=doc, image_path='assets\column_joint_demo.png')
    add_scaled_image(doc=doc, image_path='assets\column_joint_demo_2.png')

    p = doc.add_paragraph('梁撓曲鋼筋之應力假設為1.25fy')
    set_chinese_font(p, font=chinese_font)

    # Add mathematical formulas
    p = doc.add_heading('計算Ts1及Mpr-(忽略壓力筋效應)', level=3)
    set_chinese_font(p, font=chinese_font)
    p = doc.add_paragraph()
    add_omml_equation(p, [create_subscript('T', 's1'),
                          create_math_element('='),
                          create_subscript('A', 's1'),
                          create_math_element('×1.25fy')])
    run = p.add_run('，由斷面上之力平衡知')
    set_chinese_font(run, font=chinese_font)
    add_omml_equation(p, [create_subscript('C', 'c1'),
                          create_math_element('='),
                          create_subscript('T', 's1')])

    mpr_minus_equation = [create_superscript_subscript('M', '-', 'pr'),
                          create_math_element('='),
                          create_subscript('T', 's1'),
                          create_math_element('×(d-'),
                          create_fraction([create_math_element('1')],
                                          [create_math_element('2')]),
                          create_fraction([create_subscript('T', 's1')],
                                          [create_math_element('0.85'),
                                           create_superscript_subscript(
                                              'f', '\'', 'c'),
                                           create_math_element('b')]),
                          create_math_element(')')]

    p = doc.add_paragraph()
    add_omml_equation(p, mpr_minus_equation)

    p = doc.add_heading('計算Ts2及Mpr+(忽略壓力筋效應)', level=3)
    set_chinese_font(p, font=chinese_font)

    p = doc.add_paragraph()
    add_omml_equation(p, [create_subscript('T', 's2'),
                          create_math_element('='),
                          create_subscript('A', 's2'),
                          create_math_element('×1.25fy')])
    run = p.add_run('，由斷面上之力平衡知')
    set_chinese_font(run, font=chinese_font)
    add_omml_equation(p, [create_subscript('C', 'c2'),
                          create_math_element('='),
                          create_subscript('T', 's2')])

    p = doc.add_paragraph()
    mpr_plus_equation = [create_superscript_subscript('M', '+', 'pr'),
                         create_math_element('='),
                         create_subscript('T', 's2'),
                         create_math_element('×(d-'),
                         create_fraction([create_math_element('1')],
                                         [create_math_element('2')]),
                         create_fraction([create_subscript('T', 's2')],
                                         [create_math_element('0.85'),
                                         create_superscript_subscript(
                                             'f', '\'', 'c'),
                                         create_math_element('b')]),
                         create_math_element(')')]
    add_omml_equation(p, mpr_plus_equation)
    p = doc.add_paragraph('計算柱端剪力')
    set_chinese_font(p, font=chinese_font)

    p = doc.add_paragraph()
    vh_equation = [create_subscript('V', 'h'),
                   create_math_element('='),
                   create_fraction([
                       create_superscript_subscript('M', '+', 'pr'),
                       create_math_element('+'),
                       create_superscript_subscript('M', '-', 'pr')
                   ], [
                       create_fraction([
                           create_subscript('H', '1'),
                           create_math_element('+'),
                           create_subscript('H', '2')
                       ], [
                           create_math_element('2')
                       ])
                   ])]
    add_omml_equation(p, vh_equation)

    p = doc.add_paragraph('計算設計剪力')
    set_chinese_font(p, font=chinese_font)
    p = doc.add_paragraph()
    add_omml_equation(p, [create_subscript('V', 'u'),
                          create_math_element('='),
                          create_subscript('T', 's1'),
                          create_math_element('+'),
                          create_subscript('C', 'c2'),
                          create_math_element('-'),
                          create_subscript('V', 'h')])
    p = doc.add_heading('接頭剪力設計例', level=2)
    set_chinese_font(p, font=chinese_font)
    p = doc.add_paragraph('以新規範401-110檢討梁柱接頭剪力強度，其適用樣態如下表所示：')
    set_chinese_font(p, font=chinese_font)
    add_image_to_doc(doc=doc, image_path='assets\column_joint_code.png')
    count = 0
    for column in column_list:
        if not (column.serial in output_serial or column.floor in output_floor):
            continue
        for pos, result in column.joint_result.items():
            create_case_result_sheet(doc=doc,
                                     chinese_font=chinese_font,
                                     result=result,
                                     floor=column.floor,
                                     serial=column.serial,
                                     fc=column.fc,
                                     mpr_minus_equation=mpr_minus_equation,
                                     mpr_plus_equation=mpr_plus_equation,
                                     vh_equation=vh_equation
                                     )
            count += 1

    # Save the document
    doc.save(doc_filename)


def create_case_result_sheet(doc: DocumentObject,
                             chinese_font: str,
                             result: dict,
                             floor: str,
                             serial: str,
                             fc,
                             mpr_minus_equation,
                             mpr_plus_equation,
                             vh_equation):
    vh1_equation = [create_subscript('V', 'h'),
                    create_math_element('='),
                    create_fraction([
                        create_superscript_subscript('M', '+', 'pr1'),
                        create_math_element('+'),
                        create_superscript_subscript('M', '-', 'pr2')
                    ], [
                        create_fraction([
                            create_subscript('H', '1'),
                            create_math_element('+'),
                            create_subscript('H', '2')
                        ], [
                            create_math_element('2')
                        ])
                    ])]

    vh2_equation = [create_subscript('V', 'h'),
                    create_math_element('='),
                    create_fraction([
                        create_superscript_subscript('M', '-', 'pr1'),
                        create_math_element('+'),
                        create_superscript_subscript('M', '+', 'pr2')
                    ], [
                        create_fraction([
                            create_subscript('H', '1'),
                            create_math_element('+'),
                            create_subscript('H', '2')
                        ], [
                            create_math_element('2')
                        ])
                    ])]

    left_beam = result['left_beam']
    right_beam = result['right_beam']
    pos = result['pos']
    x11 = result['x11']
    x21 = result['x21']
    x12 = result['x12']
    x22 = result['x22']
    hc = result['hc']
    bj = result['bj']
    Aj = result['Aj']
    Ts1_top = result['Ts1_top']
    Ts1_bot = result['Ts1_bot']
    Ts2_top = result['Ts2_top']
    Ts2_bot = result['Ts2_bot']
    Vh1 = result['Vh1(tf)']
    Vh2 = result['Vh2(tf)']
    Vn = result['Vn(tf)']
    Vu = result['Vu(tf)']
    Mpr1_minus = result['Mpr1-(tf-m)']
    Mpr1_plus = result['Mpr1+(tf-m)']
    Mpr2_minus = result['Mpr2-(tf-m)']
    Mpr2_plus = result['Mpr2+(tf-m)']
    H1 = result['H1']/100
    H2 = result['H2']/100
    design_code = result['design_code']
    _code_15_2_6 = result['_code_15_2_6']
    _code_15_2_7 = result['_code_15_2_7']
    _code_15_2_8 = result['_code_15_2_8']
    dcr = result['DCR']
    message = ""
    pass_sign = ""
    if dcr <= 1:
        pass_sign = "≥"
        message = "→OK，檢核通過"
    else:
        pass_sign = "<"
        message = "→NG!"

    p = doc.add_heading(f'計算{floor}-{serial} {pos}向剪力強度', level=4)
    set_chinese_font(p, font=chinese_font)

    # Create a table
    # Set table background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9EAD3')

    code_table = doc.add_table(rows=1, cols=2, style='Table Grid')
    code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    code_table_cells = code_table.rows[0].cells
    code_table_cells[0].text = "接頭性質"
    code_table_cells[1].text = "是否符合"
    codes = [(["柱連續性"], [_code_15_2_6]),
             (["梁連續性"], [_code_15_2_7]),
             (["橫向梁圍束"], [_code_15_2_8]),
             (["依以上內容判斷，", [create_subscript('V', 'n'),
                            create_math_element(f"={design_code}λ"),
                            create_square_root(
                                create_superscript_subscript('f', '\'', 'c')),
                            create_subscript('A', 'j')]], [])]
    for items, values in codes:
        row_cells = code_table.add_row().cells
        p_item = row_cells[0].paragraphs[0]
        p_value = row_cells[1].paragraphs[0]

        # run_item = p_item.add_run(item)
        p_item.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if not values:
            new_cells = row_cells[0].merge(row_cells[-1])
            p_item = new_cells.paragraphs[0]
        if not items:
            new_cells = row_cells[0].merge(row_cells[-1])
            p_value = new_cells.paragraphs[0]
        for item in items:
            if isinstance(item, str):
                run = p_item.add_run(item)
                set_chinese_font(run, chinese_font)
            elif isinstance(item, list):
                add_omml_equation(p_item, item)

        for value in values:
            if isinstance(value, str):
                run = p_value.add_run(value)
                set_chinese_font(run, chinese_font)
            elif isinstance(value, list):
                add_omml_equation(p_value, value)

        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    # Create a table
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table.rows[0].cells:
        cell._element.get_or_add_tcPr().append(deepcopy(shading_elm))

    for cell in code_table.rows[0].cells:
        cell._element.get_or_add_tcPr().append(deepcopy(shading_elm))

    # Header row formatting
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '項目'
    hdr_cells[1].text = '數值'
    for cell in hdr_cells:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        set_chinese_font(run, chinese_font)

    # Add rows to the table with formulas
    rows = [(['平行剪力方向X向柱深', [create_subscript('h', 'c')]], f'{hc} cm'),
            (['左梁'], left_beam),
            (['右梁'], right_beam)]
    if left_beam != "X":
        rows.extend([
            (['左梁邊與柱邊距離', [create_subscript('x', '11')]],
             f'min({x11} cm，{hc}/4) = {min(x11,hc/4)} cm'),
            (['左梁邊與柱邊距離', [create_subscript('x', '12')]],
             f'min({x12} cm，{hc}/4) = {min(x12,hc/4)} cm')
        ])
    if right_beam != "X":
        rows.extend([
            (['右梁邊與柱邊距離', [create_subscript('x', '21')]],
             f'min({x21} cm，{hc}/4) = {min(x22,hc/4)} cm'),
            (['右梁邊與柱邊距離', [create_subscript('x', '22')]],
             f'min({x22} cm，{hc}/4) = {min(x22,hc/4)} cm')
        ])

    rows.extend([(['接頭有效抗剪寬度', [create_subscript('b', 'j'),
                                create_math_element('='),
                                create_subscript('b', 'b'),
                                create_math_element('+'),
                                create_subscript('x', '1'),
                                create_math_element('+'),
                                create_subscript('x', '2')],
                   '若左右接有梁則取兩者平均'], f'{bj} cm'),
                 (['有效抗剪水平斷面積', [create_subscript('A', 'j'),
                                 create_math_element('='),
                                 create_subscript('b', 'j'),
                                 create_math_element('×'),
                                 create_subscript('h', 'c')]], f'{Aj} cm²')])
    if left_beam:
        rows.append(
            (['左梁上層受拉', [create_subscript('T', 's1'), "=",
                         create_subscript('C', 'c1'), "=",
                         create_subscript('A', 's1,top'), "×1.25fy"]], f'{Ts1_top} tf')
        )
        rows.append(([[create_superscript_subscript('M', '-', 'pr1'), "="],
                     deepcopy(mpr_minus_equation)], f'{Mpr1_minus} tf-m'))
        rows.append(
            (['左梁下層受拉', [create_subscript('T', 's2'), "=",
                         create_subscript('C', 'c2'), "=",
                         create_subscript('A', 's1,bot'), "×1.25fy"]], f'{Ts1_bot} tf')
        )
        rows.append(([[create_superscript_subscript('M', '+', 'pr1'), "="],
                    deepcopy(mpr_plus_equation)], f'{Mpr1_plus} tf-m'))
    if right_beam:
        rows.append(
            (['右梁上層受拉', [create_subscript('T', 's1'), "=",
                         create_subscript('C', 'c1'), "=",
                         create_subscript('A', 's2,top'), "×1.25fy"]], f'{Ts2_top:.2f} tf')
        )
        rows.append(([[create_superscript_subscript('M', '-', 'pr2'), "="],
                      deepcopy(mpr_minus_equation)], f'{Mpr2_minus:.2f} tf-m'))
        rows.append(
            (['右梁下層受拉', [create_subscript('T', 's2'), "=",
                         create_subscript('C', 'c2'), "=",
                         create_subscript('A', 's2,bot'), "×1.25fy"]], f'{Ts2_bot:.2f} tf')
        )
        rows.append(([[create_superscript_subscript('M', '+', 'pr2'), "="],
                    deepcopy(mpr_plus_equation)], f'{Mpr2_plus:.2f} tf-m'))
    rows.extend([
        ([[create_fraction([
            create_subscript('H', '1'),
            create_math_element('+'),
            create_subscript('H', '2')
        ], [
            create_math_element('2')
        ])
        ]], f'({H1} + {H2}) / 2 ={(H1+H2)/2} m'),
        ([[create_subscript('V', 'h1'), "="],
          deepcopy(vh1_equation)], f'({Mpr1_plus:.2f} + {Mpr2_minus:.2f}) / {(H1+H2)/2}={Vh1:.2f} tf'),
        ([[create_subscript('V', 'h2'), "="],
          deepcopy(vh2_equation)], f'({Mpr1_minus:.2f} + {Mpr2_plus:.2f}) / {(H1+H2)/2}={Vh2:.2f} tf'),
        ([[create_subscript('V', 'u'), "=",
           create_subscript('T', 's1'), "+",
           create_subscript('C', 'c2'), "-",
           create_subscript('V', 'h')]], f'{Vu:.2f} tf'),
        (['剪力強度',
          [create_math_element('∅'),
           create_subscript('V', 'n×'), "=",
           create_math_element(f'0.85×{design_code}λ'),
           create_square_root(create_superscript_subscript('f', '\'', 'c')),
           create_subscript('A', 'j')]], f'0.85×{design_code}×\u221A{fc}×{Aj}={Vn:.2f} tf'),
        ([[create_math_element('∅'),
           create_subscript('V', 'n×'),
           create_math_element(pass_sign),
           create_subscript('V', 'u')], message], [])
    ])

    for items, values in rows:
        row_cells = table.add_row().cells
        p_item = row_cells[0].paragraphs[0]
        p_value = row_cells[1].paragraphs[0]

        # run_item = p_item.add_run(item)
        p_item.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if not values:
            new_cells = row_cells[0].merge(row_cells[-1])
            p_item = new_cells.paragraphs[0]
        if not items:
            new_cells = row_cells[0].merge(row_cells[-1])
            p_value = new_cells.paragraphs[0]
        for item in items:
            if isinstance(item, str):
                run = p_item.add_run(item)
                set_chinese_font(run, chinese_font)
            elif isinstance(item, list):
                add_omml_equation(p_item, item)

        for value in values:
            if isinstance(value, str):
                run = p_value.add_run(value)
                set_chinese_font(run, chinese_font)
            elif isinstance(value, list):
                add_omml_equation(p_value, value)

        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                            bottom={"sz": 12, "val": "single",
                                    "color": "000000"},
                            left={"sz": 12, "val": "single", "color": "000000"},
                            right={"sz": 12, "val": "single", "color": "000000"})


if __name__ == "__main__":
    create_calculate_sheet('test.docx', [])
